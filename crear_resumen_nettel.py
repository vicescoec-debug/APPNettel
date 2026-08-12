from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"C:\Users\espin\Documents\APP Nettel\Resumen_Tecnico_APP_Nettel_Maritimo.docx"
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 100, 112)

def font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd")) or OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc.get_or_add_tcPr()
    tcMar = tc.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tc.append(tcMar)
    for tag, value in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node = tcMar.find(qn(f"w:{tag}")) or OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")
        tcMar.append(node)

def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    ind = OxmlElement("w:tblInd"); ind.set(qn("w:w"), "120"); ind.set(qn("w:type"), "dxa"); tblPr.append(ind)
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i]/1440)
            tcW = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tcW.set(qn("w:w"), str(widths[i])); tcW.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cell)

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.5)
    p.paragraph_format.first_line_indent = Inches(-.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    font(p.add_run(text))
    return p

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.right_margin = sec.bottom_margin = sec.left_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in (("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",12,DARK,8,4)):
    s=styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=color
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(header.add_run("NETTEL MARÍTIMO  |  INFORME TÉCNICO"), 8.5, True, MUTED)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(footer.add_run("Documento de trabajo · 22 de junio de 2026"), 8.5, False, MUTED)

p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4)
font(p.add_run("RESUMEN TÉCNICO"),23,True,RGBColor(0,0,0))
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(16)
font(p.add_run("Modernización y depuración de la aplicación Android Nettel Marítimo"),14,False,MUTED)
for label,value in (("Proyecto","Nettel Marítimo GPS para Android"),("Fecha","22 de junio de 2026"),("Estado","Compilación modernizada y APK de depuración generado"),("Versión heredada","4.0.7 · SDK 28 · código Java de 2019")):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    font(p.add_run(label+": "),11,True); font(p.add_run(value),11)

doc.add_heading("Resultado ejecutivo", level=1)
p=doc.add_paragraph()
font(p.add_run("La aplicación fue recuperada y modernizada hasta obtener una compilación correcta para Android API 35. "),11,True,DARK)
font(p.add_run("Se generó un APK instalable, se ejecutaron las pruebas unitarias existentes y se corrigieron bloqueos de compilación, compatibilidad y seguridad. La validación operativa completa todavía requiere un dispositivo Android y credenciales válidas del backend."))

doc.add_heading("Situación encontrada", level=1)
for item in [
    "Proyecto Android Java creado en 2019, configurado originalmente con SDK 28, Gradle 5.1.1 y Android Support Libraries.",
    "Configuración local vinculada a rutas de una computadora Mac anterior, sin un entorno reproducible en Windows.",
    "Dependencias retiradas de sus repositorios y recursos XML inválidos que impedían compilar.",
    "Servicios SOAP y PHP mediante HTTP sin cifrado, además de una clave de Google Maps incluida en los recursos.",
    "Servicio de alertas incompatible con las restricciones de ejecución y notificaciones de Android moderno.",
    "Registro de usuario y contraseña en consola y consultas SQLite construidas mediante concatenación de texto."
]: add_bullet(doc,item)

doc.add_heading("Trabajo realizado", level=1)
rows=[
    ("Entorno","JDK 17 y Android SDK portátil dentro del espacio de trabajo."),
    ("Compilación","Gradle 8.9, Android Gradle Plugin 8.7.3, compileSdk/targetSdk 35."),
    ("Dependencias","Corrección de Volley y sustitución de bibliotecas desaparecidas por controles Android nativos."),
    ("Manifest","Componentes exportados explícitamente, permisos modernos y servicio foreground de sincronización."),
    ("Alertas","Canales de notificación, PendingIntent inmutables y permiso POST_NOTIFICATIONS."),
    ("Red","Tráfico HTTP bloqueado por defecto y permitido temporalmente solo para los servidores heredados."),
    ("Seguridad","Eliminación del registro de contraseñas y reducción de exposición de componentes."),
    ("Datos","Inserciones SQLite con ContentValues y consulta de clientes mediante parámetros.")]
table=doc.add_table(rows=1,cols=2); table.style="Table Grid"
set_table_geometry(table,[2700,6660])
for i,t in enumerate(("Área","Mejora aplicada")):
    c=table.rows[0].cells[i]; shade(c,"F2F4F7"); font(c.paragraphs[0].add_run(t),10.5,True,DARK)
for a,b in rows:
    cells=table.add_row().cells
    font(cells[0].paragraphs[0].add_run(a),10.5,True)
    font(cells[1].paragraphs[0].add_run(b),10.5)
set_table_geometry(table,[2700,6660])

doc.add_heading("Validaciones realizadas", level=1)
for item in [
    "Tarea assembleDebug completada correctamente.",
    "Pruebas unitarias: 1 ejecutada, 1 aprobada, sin errores.",
    "APK de depuración generado correctamente, con tamaño aproximado de 4,6 MB.",
    "Los endpoints heredados respondieron durante la revisión; no se validó una sesión autenticada real."
]: add_bullet(doc,item)

doc.add_heading("Entregables", level=1)
table=doc.add_table(rows=1,cols=3); table.style="Table Grid"
set_table_geometry(table,[1850,2450,5060])
for i,t in enumerate(("Elemento","Estado","Ubicación")):
    c=table.rows[0].cells[i]; shade(c,"F2F4F7"); font(c.paragraphs[0].add_run(t),10.5,True,DARK)
for vals in [
    ("Proyecto Android","Actualizado","APP Nettel/Nettel Maritimo/Android/GPS/GPS"),
    ("APK debug","Generado","app/build/outputs/apk/debug/app-debug.apk"),
    ("Entorno portátil","Configurado","APP Nettel/.tools")]:
    cells=table.add_row().cells
    for i,v in enumerate(vals): font(cells[i].paragraphs[0].add_run(v),10)
set_table_geometry(table,[1850,2450,5060])

doc.add_heading("Pendientes para puesta en producción", level=1)
for item in [
    "Instalar el APK en un dispositivo o emulador y probar Android 10 a 15.",
    "Validar login, recuperación de contraseña, mapas, dispositivos, históricos y alertas con credenciales reales.",
    "Migrar los servicios SOAP/PHP a HTTPS y retirar la excepción temporal de tráfico HTTP.",
    "Revocar y reemplazar la clave de Google Maps expuesta; restringir la nueva clave por paquete y huella SHA-1.",
    "Confirmar la estrategia de alertas en segundo plano. El servicio foreground funciona, pero Android 15 limita las tareas prolongadas de sincronización.",
    "Generar un Android App Bundle release firmado, gestionar la firma fuera del repositorio y realizar pruebas de publicación.",
    "Ampliar la cobertura automática; la prueba existente solo verifica una suma y no cubre la lógica real de negocio."
]: add_bullet(doc,item)

doc.add_heading("Conclusión", level=1)
p=doc.add_paragraph()
font(p.add_run("El proyecto ya no está bloqueado técnicamente: compila, produce un APK y cuenta con una base compatible con Android actual. "),11,True,DARK)
font(p.add_run("El siguiente hito debe ser una prueba funcional controlada contra el backend real. Después de esa validación se podrán corregir incidencias de negocio, completar la migración HTTPS y preparar una versión release para distribución."))

doc.core_properties.title = "Resumen técnico - APP Nettel Marítimo"
doc.core_properties.subject = "Modernización y depuración de aplicación Android"
doc.core_properties.author = "Equipo de proyecto Nettel"
doc.save(OUT)
print(OUT)
